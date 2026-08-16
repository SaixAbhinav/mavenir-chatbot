import docx
import pytest


@pytest.fixture
def make_docx(tmp_path):
    """Build a .docx from (style, content) pairs.

    Styles: 'Heading 1'..'Heading 4', 'Normal' — content is paragraph text.
    Style 'Table' — content is a list of rows, each row a list of cell strings;
    the first row becomes the Markdown header row.
    """
    def _make(name, blocks):
        document = docx.Document()
        for style, content in blocks:
            if style == "Table":
                rows = content
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        table.cell(r, c).text = cell_text
            else:
                document.add_paragraph(content, style=style)
        path = tmp_path / name
        document.save(path)
        return path
    return _make
